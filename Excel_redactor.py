import collections
import datetime
import math
import os
import shutil
import time
from copy import copy
import binascii

# pip install docxcompose

# import locale
# os.environ["PYTHONIOENCODING"] = "utf-8";
# myLocale=locale.setlocale(category=locale.LC_ALL, locale="en_GB.UTF-8")

import docx
import openpyxl
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt
from docx.table import _Cell
from docxcompose.composer import Composer
from openpyxl.styles import Border, Side, Alignment, Font
from openpyxl import Workbook, load_workbook
from openpyxl.worksheet.dimensions import ColumnDimension, DimensionHolder
from openpyxl.utils import get_column_letter
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from decimal import Decimal
from PyQt6.QtCore import QThread, pyqtSignal, QObject, QRunnable, pyqtSlot

import traceback, sys
def change_gas(cur_list):
    cur_str = str(cur_list)
    pos = cur_str.find(" (")
    if pos != -1:
        new_str = cur_str[:pos]
    else:
        new_str = cur_list
    return new_str

def get_super(x):
    normal = "0123456789+-=()"
    super_s = "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾"
    res = x.maketrans("".join(normal), "".join(super_s))
    return x.translate(res)

def isDigit(x):
    try:
        float(x)
        return True
    except ValueError:
        return False

def change_state_with_old(r_otv, n_gas, old_state):
    r_otv = r_otv.strip()

    if 'ПР' in old_state:
        res = 'ПР'
    elif 'ГР' in old_state:
        res = 'ГР'
    else:
        # если у нас число - считаем
        # если пусто или - или текст
        if isDigit(r_otv):  # if otverstie est' -
            r_otv = float(r_otv)
            d_otv = math.sqrt((r_otv * 4) / math.pi) * 1000
            round_num = round(d_otv * 2) / 2
            if str(round_num)[-1] == "0":
                res = str(round_num)
                pos = res.find(".0")
                res = res[:pos]
                res += " мм"
                # print(res)
            else:
                res = str(round_num)
                res = res.replace(".", ",")
                res += " мм"
        else:
            res = 'ПР'

    if n_gas is not None and n_gas != "None":
        new_state = n_gas + ' ' + res
    else:
        new_state = 'Неопределенное вещество ' + res

    return new_state

def another_e(chislo):
    power = int(str(chislo)[-2]+str(chislo)[-1])
    value_new = float(str(chislo * pow(10, power))[:-2])
    answer = str(round(value_new, 2))
    p = get_super(str(10))
    answer = answer + ' × 10' + p
    return answer

def filtery(sort_dict, obor, ov, pdo, max_1,max_2):
    #Структура: [0]Оборудование, [1]Состояние, [2]Опасное вещество,
    # [3]Площадь отверстия, [4]Масса, [5]Дрейф, [6]1,[7]10,[8]25,[9]50,[10]90,[11]99
    rez = []
    #1.Выбираем ооборудование
    #2.Выбираем опасное вещество
    for o in range(len(obor)):
        for v in range(len(ov)):
            for p in range(len(pdo)):
                a = []
                max = 0
                max_dr = 0
                b = []
                max_str = []
                for i in range(len(sort_dict)):
                    if sort_dict[i][0] == obor[o] and sort_dict[i][2] == ov[v] and sort_dict[i][3] == pdo[p]:
                        a.append(sort_dict[i])
                    if i == len(sort_dict)-1:
                        for j in range(len(a)):
                            if a[j][max_1] is None: a[j][max_1] = 0
                            if a[j][max_1] > max:
                                b.clear()
                                b.append(a[j])
                                max = a[j][max_1]
                            elif a[j][max_1] == max:
                                b.append(a[j])
                        for j in range(len(b)):
                            if b[j][max_2] is None: b[j][max_2] = 0
                            if b[j][max_2] >= max_dr:
                                max_str = (b[j])
                                max_dr = b[j][max_2]
                        if max_str is not []:
                            rez.append(max_str)
    rez = [x for x in rez if x != []]
    # print('1111')
    # print(rez)
    for i in range (len(rez)):
        rez[i].pop(3)
        rez[i].pop(2)

    # print(rez)
    # print('1111')
    return rez

def filtery_for_3_scens(sort_dict, obor, ov, ishod,  pdo):
    #print(pdo)
    #print(sort_dict)
    rez = []
    for o in range(len(obor)):
        for v in range(len(ov)):
            for ish in range(len(ishod)):
                for p in range(len(pdo)):
                    sum_chast = 0
                    a = []
                    for i in range(len(sort_dict)):
                        if sort_dict[i][0] == obor[o] and sort_dict[i][1] == ov[v] and sort_dict[i][2] == ishod[ish] and sort_dict[i][3] == pdo[p]:
                            a.append(sort_dict[i])
                        if i == len(sort_dict)-1:
                            for j in range(len(a)):
                                 sum_chast = sum_chast + a[j][4]
                            if sum_chast != 0:
                                a[j][4] = sum_chast
                                rez.append(a[j])
    rez = [x for x in rez if x != []]
    return rez

def filtery_for_3(sort_dict, obor, ov, pdo, max_1,max_2,max_3):
    #Структура: [0]Оборудование, [1]Состояние, [2]Опасное вещество,
    # [3]Площадь отверстия, [4]Масса, [5]Дрейф, [6]1,[7]10,[8]25,[9]50,[10]90,[11]99
    rez = []
    for o in range(len(obor)):
        for v in range(len(ov)):
            for p in range(len(pdo)):
                a = []
                max = 0
                max_dr = 0
                max_dr_dr = 0
                b = []
                c = []
                max_str = []
                for i in range(len(sort_dict)):
                    if sort_dict[i][0] == obor[o] and sort_dict[i][2] == ov[v] and sort_dict[i][3] == pdo[p]:
                        a.append(sort_dict[i])
                    if i == len(sort_dict)-1:
                        for j in range(len(a)):
                            if a[j][max_1] is None: a[j][max_1] = 0
                            if a[j][max_1] > max:
                                b.clear()
                                b.append(a[j])
                                max = a[j][max_1]
                            elif a[j][max_1] == max:
                                b.append(a[j])
                        for j in range(len(b)):
                            if b[j][max_2] is None: b[j][max_2] = 0
                            if b[j][max_2] > max_dr:
                                c.clear()
                                c.append(b[j])
                                max_dr = b[j][max_2]
                            elif b[j][max_2] == max_dr:
                                c.append(b[j])
                        for j in range(len(c)):
                            if c[j][max_3] is None or c[j][max_3] == 'None': c[j][max_3] = 0
                            if c[j][max_3] > max_dr_dr:
                                max_str = (c[j])
                                max_dr_dr = c[j][max_2]
                        if max_str is not []:
                            rez.append(max_str)
    rez = [x for x in rez if x != []]
    for i in range (len(rez)):
        rez[i].pop(3)
        rez[i].pop(2)
    return rez

def filtery_for_1(sort_dict, obor, ov, pdo, max_1):
    #Структура: [0]Оборудование, [1]Состояние, [2]Опасное вещество,
    # [3]Площадь отверстия, [4]Масса, [5]Дрейф, [6]1,[7]10,[8]25,[9]50,[10]90,[11]99
    rez = []
    for o in range(len(obor)):
        for v in range(len(ov)):
            for p in range(len(pdo)):
                a = []
                max = 0
                max_str = []
                for i in range(len(sort_dict)):
                    if sort_dict[i][0] == obor[o] and sort_dict[i][2] == ov[v] and sort_dict[i][3] == pdo[p]:
                        a.append(sort_dict[i])
                    if i == len(sort_dict)-1:
                        for j in range(len(a)):
                            if a[j][max_1] is None: a[j][max_1] = 0
                            if a[j][max_1] >= max:
                                max_str = (a[j])
                                max = a[j][max_1]
                        if max_str is not []:
                            rez.append(max_str)
    rez = [x for x in rez if x != []]
    for i in range (len(rez)):
        rez[i].pop(3)
        rez[i].pop(2)
    return rez

def filtery_for_vert_li_fakel(sort_dict, obor, ov, pdo, max_1):
    rez = []
    for o in range(len(obor)):
        for v in range(len(ov)):
            for p in range(len(pdo)):
                a = []
                max = 0
                max_str = []
                for i in range(len(sort_dict)):
                    if sort_dict[i][0] == obor[o] and sort_dict[i][2] == ov[v] and sort_dict[i][3] == pdo[p]:
                        a.append(sort_dict[i])
                    if i == len(sort_dict)-1:
                        for j in range(len(a)):
                            if a[j][max_1] is None: a[j][max_1] = 0
                            if a[j][max_1] >= max:
                                max_str = (a[j])
                                max = a[j][max_1]
                        if max_str is not []:
                            rez.append(max_str)
    rez = [x for x in rez if x != []]
    return rez

def chistka_3_ogn(dict):
    for i in range (len(dict)):
        dict[i].pop(3)
        dict[i].pop(2)
    return dict

def chistka_for_pozh_intens(dict):
    for i in range (len(dict)):
        dict[i].pop(3)
        dict[i].pop(1)
    return dict

def Delete_table(table, doc):
    #print(doc.tables[table]._element.getparent())
    #delete_paragraph(doc.paragraphs[table-1])
    doc.paragraphs[table*2-2].clear()
    doc.tables[table]._element.getparent().remove(doc.tables[table]._element)

def chistka_vert_fakel(dict):
    for i in range (len(dict)):
        dict[i].pop(3)
        dict[i].pop(1)
    return dict

def combine_all_docx(filename_master, files_list, file_name):
    number_of_sections=len(files_list)
    master = Document(filename_master)
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document(files_list[i])
        composer.append(doc_temp)
    composer.save(file_name)

    doc = Document(file_name)
    n_record = 1

    for y in range(len(doc.paragraphs)):
        print(len(doc.paragraphs))
        print(doc.paragraphs)
        if doc.paragraphs[y].text != " " and doc.paragraphs[y].text != "" and doc.paragraphs[y].text is not None:
        # if y % 2 == 0:
            doc.paragraphs[y].text = "Таблица " + str(n_record) + ". " + doc.paragraphs[y].text
            n_record = n_record + 1
    doc.save(file_name)

def start_combine(file_name):
    folder = os.listdir('output')
    print(folder)
    file_master = "output/" + folder[0]
    print(file_master)
    files = []
    if len(folder) == 1:
        shutil.copyfile(file_master, file_name)
    if len(folder) == 2:
        files.append("output/" + folder[1])
        combine_all_docx(file_master, files, file_name)
    if len(folder) > 2:
        for x in range(len(folder) - 1):
            files.append("output/" + folder[x + 1])
        combine_all_docx(file_master, files, file_name)

class Worker_Excel(QObject):

    _signal = pyqtSignal(int)
    finished = pyqtSignal(int)
    error = pyqtSignal(int)
    # terminate = pyqtSignal()

    def __init__(self, excel_path, word_path):
        super(Worker_Excel, self).__init__()
        self.excel_path = excel_path
        self.word_path = word_path
    #
    # def __del__(self):
    #     self.wait()
    def run(self):
        print("IN THE RUN")
        start_time = time.time()

        wb = load_workbook(self.excel_path)

        sheet1 = wb.get_sheet_names()
        sheet = wb.get_sheet_by_name(sheet1[0])


        if 'Scens' in sheet1[0]:  # RPR, i = 6


            # B - Наименование оборудования
            # D - опасное вещество
            # Е - исход
            # G - диаметр !!!!!!!!!!!!!!!!!!!!
            # F - частота сценария
            # H - масса Гф
            # I - масса Жф

            j = 0
            dict = []
            l = 0

            all_rows = sheet.max_row

            oborudovanie = []
            ov = []
            pdo = []
            ishod = []

            for i in range(3, all_rows):
                stroka = []
                obor = str(sheet['B' + str(i)].value).strip()  # Оборудование
                if not (obor in oborudovanie): oborudovanie.append(obor)
                stroka.append(obor)

                ve = (sheet['D' + str(i)].value)  # ОВ
                if not (ve in ov): ov.append(ve)
                stroka.append(ve)

                ish = (sheet['E' + str(i)].value)  # Исход
                if not (ish in ishod): ishod.append(ish)
                stroka.append(ish)

                q = str(sheet['G' + str(i)].value).strip()  # ПДО
                if q == '-':
                    q = 0
                    stroka.append(q)
                else:
                    if q is None or q == 'None':
                        q = 0
                        stroka.append(q)
                    else:
                        q = (math.sqrt(float(q) * 4 / float(math.pi)) * 1000)
                        # print(q)
                        if int(q) == 12 or int(q) == 13:
                            q = 12.5
                            stroka.append(q)

                        else:
                            q = int(q)
                            stroka.append(q)
                if not (q in pdo): pdo.append(q)

                with_10_ = (sheet['F' + str(i)].value)  # Частота сценария
                stroka.append(with_10_)

                stroka.append(sheet['H' + str(i)].value)  # Масса ГФ
                stroka.append(sheet['I' + str(i)].value)  # Масса Жф

                dict.append(stroka)

            # print(dict) # [Оборудование[0], ОВ[1], Исход[2], ПДО[3], Частота сценария!![4], Масса Гф, Масса ЖФ]

            dict.sort(key=lambda row: (row[1] or 0, row[2] or 0, row[3] or 0, row[0] or 0))
            print('1')
            rez = filtery_for_3_scens(dict, oborudovanie, ov, ishod, pdo)
            # print(rez)

            for p in (range(len(rez))):
                rez[p][4] = another_e(rez[p][4])

            # print(rez)

            doc = Document('Templates/Scene.docx')
            doc_table = doc.tables

            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(10)
            paragraph_format = doc.styles['Normal'].paragraph_format
            paragraph_format.line_spacing = Pt(12)  ###################
            paragraph_format.space_after = (0)  ################

            step = len(rez) / 100
            percent = 1
            print('2')

            for k in range(len(rez)):

                if 1.0 > k % step >= 0.0:
                    self._signal.emit(1)

                list = rez[k]
                row = doc_table[0].add_row().cells  # олучаем все ячейки ряда
                for col in range(7):
                    cell = row[col]
                    paragraph = cell.paragraphs[0]
                    paragraph.style = doc.styles['Normal']
                    if col == 3:
                        if (list[col]) == 0:
                            list[col] = 'Полное разрущение'
                        else:
                            list[col] = str(list[col]).replace('.', ',')
                    if col == 4: list[col] = str(list[col]).replace('.', ',')
                    if col == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(str(list[col]))

            print("SAVE------------------------")
            doc.save(self.word_path)
            print("--- %s seconds ---" % (time.time() - start_time))
            self.finished.emit(0)
        #return 0

        else:
            wb.close()
            print("WRONG STRUCT")
            self.error.emit(1)
            self.finished.emit(-1)
            return 1

        #
        # j = 0
        # step = all_rows / 100
        # percent = 1
        # #self._signal.emit(percent)
        # for row in range(2, all_rows-2):
        #     # print(row)
        #     # print("%%%%%%%%%%%% -- " + str(percent))
        #     # print(complete_percent_value)
        #     if row % step < 1.0 and row % step >= 0.0:
        #         #percent+=percent
        #         percent += 1
        #         self._signal.emit(percent)
        #     list = dict[j]
        #     j = j + 1
        #     row = doc_table[0].add_row().cells  # олучаем все ячейки ряда
        #     for col in range(7):
        #         cell = row[col]
        #         paragraph = cell.paragraphs[0]
        #         paragraph.style = doc.styles['Normal']
        #         if col == 3:
        #             if (list[col]) == 0:
        #                 list[col] = 'Полное разрущение'
        #             else:
        #                 list[col] = str(list[col]).replace('.', ',')
        #         if col == 4: list[col] = str(list[col]).replace('.', ',')
        #         if col == 0:
        #             paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #         else:
        #             paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #         run = paragraph.add_run(str(list[col]))
        #
        #         cell = doc_table[0].cell(row, col)
        #         cell.text = str(list[col])
        #
        #
        # pos = excel_path.find('.xls')
        # print("I AM HERE")
        # print(pos)
        # copy_path = excel_path
        # copy_path = excel_path[:pos] + " -copy" + excel_path[pos:]
        # print(copy_path)
        # print(excel_path)
        #
        # shutil.copy(excel_path, copy_path)
        # book = openpyxl.load_workbook(copy_path)
        # cur_sheet = book.active
        # value = 0
        #
        #
        # for row in cur_sheet[3:cur_sheet.max_row]:
        #     # print('in loop')
        #     cur_cell = row[6]
        #     if cur_cell.value == '-':
        #         # print('Complete ')
        #         continue
        #     else:
        #         if math.sqrt(float(cur_cell.value) * 4 / float(math.pi))*1000 < 0 or None:
        #             print('BO4ik potik')
        #         else:
        #             value = math.sqrt(float(cur_cell.value) * 4 / float(math.pi))*1000
        #
        #         # #print(value)
        #         # print('in loop 2')
        #
        # value = str(value)
        #
        # value = value.replace('.', ',')
        # print("VALUE , = " + value)
        # index = value.find(',')
        # value = value[:index+3]
        # print("VALUE = " + value)
        #     #cell.value = '=КОРЕНЬ(F' + str(myiter) + '*4/ПИ())*1000'
        #
        # book.save(copy_path)
        # book = openpyxl.load_workbook(copy_path, data_only=True)
        #
        # cur_sheet = book.active
        # cur_cell = cur_sheet['F3']
        #
        #
        #
        # print(power)
        # print(value_new)
        # #print(value * pow(10, power))
        #
        # #print(str(value)[-1])
        # os.remove(copy_path)
        #
        # book = openpyxl.load_workbook(excel_path)

class Worker_Excel_RPR(QObject):

    _signal = pyqtSignal(int)
    finished = pyqtSignal(int)
    error = pyqtSignal(int)

        # terminate = pyqtSignal()

    def __init__(self, excel_path, word_path):
        super(Worker_Excel_RPR, self).__init__()
        self.excel_path = excel_path
        self.word_path = word_path

        #
        # def __del__(self):
        #     self.wait()
    def run(self):
        start_time = time.time()
        wb = openpyxl.load_workbook(self.excel_path)
        sheet_name = (wb.get_sheet_names())
        print(sheet_name)

        percent = 100 / len(sheet_name)


        dir = os.getcwd() + "/output"

        #dir = '/output'
        for f in os.listdir(dir):
            os.remove(os.path.join(dir, f))

        for i in range(len(sheet_name)):
            if 'Пожар-вспышка по РБ РТН' in sheet_name[i]:
                var = 'OPO'
                print("OPO")
            elif 'Пожар-вспышка (методика М' in sheet_name[i]:
                var = 'RPR'
                print("RPR")
        if var != 'OPO' and var != 'RPR':
            wb.close()
            #print("WRONG STRUCT")
            self.error.emit(2)
            self.finished.emit(-1)
            return 1

        count_done_list = 0

        for i in range(len(sheet_name)):
            print("3")
            if 'Пожар пролива Вероятностн' in sheet_name[i]:  # RPR, i = 6

                if var == 'RPR':
                    doc = Document(
                        'Templates/Результаты расчета зон действия поражающих факторов при пожаре-пролива РПР.docx')
                elif var == 'OPO':
                    doc = Document(
                        'Templates/Максимальные зоны смертельного поражения при реализации аварий с пожаром пролива.docx')

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []

                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    if not (obor in oborudovanie): oborudovanie.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)

                    s_def_otv = str(sheet['H' + str(j)].value)  # Площадь деф отв)

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)
                    if not (ve in ov): ov.append(ve)

                    stroka.append(s_def_otv)
                    if not (s_def_otv in pdo): pdo.append(s_def_otv)
                    stroka.append(sheet['I' + str(j)].value)  # Масса
                    stroka.append(sheet['K' + str(j)].value)  # Радиус 1
                    stroka.append(sheet['L' + str(j)].value)  # Радиус 10
                    stroka.append(sheet['M' + str(j)].value)  # Радиус 25
                    stroka.append(sheet['N' + str(j)].value)  # Радиус 50
                    stroka.append(sheet['O' + str(j)].value)  # Радиус 90
                    stroka.append(sheet['P' + str(j)].value)  # Радиус 99
                    stroka.append(sheet['Q' + str(j)].value)  # Радиус 100
                    dict.append(stroka)

                dict.sort(key=lambda row: (row[2] or 0, row[3] or 0, row[0] or 0))

                # print(dict)
                rez = filtery(dict, oborudovanie, ov, pdo, 6, 5)

                count_stolb = 10

                # if var == 'RPR':
                #     exist_table['0'] = 1
                # elif var == 'OPO':
                #     exist_table['1'] = 1

                for k in range(len(rez)):
                    list = rez[k]
                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда
                    # if var == 'RPR':
                    #     row = doc_table[0].add_row().cells  # олучаем все ячейки ряда
                    # elif var == 'OPO':
                    #     row = doc_table[1].add_row().cells  # олучаем все ячейки ряда
                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))

                doc.save('output/0.docx')
                print("--- %s seconds ---" % (time.time() - start_time))
                print("----------------------------")
                print(len(sheet_name))

                self._signal.emit(int(percent))
                #percent += percent
                print("----------------------------")

            elif 'Взрыв ТВС Вероятностное с' in sheet_name[i]:  # i = 5

                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []

                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    if not (obor in oborudovanie): oborudovanie.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)
                    s_def_otv = str(sheet['H' + str(j)].value)  # Площадь деф отв)

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)
                    if not (ve in ov): ov.append(ve)

                    stroka.append(s_def_otv)
                    if not (s_def_otv in pdo): pdo.append(s_def_otv)
                    stroka.append(sheet['I' + str(j)].value)  # Масса
                    stroka.append(sheet['J' + str(j)].value)  # Дрейф
                    stroka.append(sheet['K' + str(j)].value)  # Радиус 1
                    stroka.append(sheet['L' + str(j)].value)  # Радиус 10
                    stroka.append(sheet['M' + str(j)].value)  # Радиус 25
                    stroka.append(sheet['N' + str(j)].value)  # Радиус 50
                    stroka.append(sheet['O' + str(j)].value)  # Радиус 90
                    stroka.append(sheet['P' + str(j)].value)  # Радиус 99

                    dict.append(stroka)

                dict.sort(key=lambda row: (row[2] or 0, row[3] or 0, row[0] or 0))
                # print(dict)
                rez = filtery(dict, oborudovanie, ov, pdo, 6, 5)

                count_stolb = 10

                if var == 'RPR':
                    doc = Document(
                        'Templates/Результаты расчета зон действия поражающих факторов при взрыве ТВС.docx')  ####################
                elif var == 'OPO':
                    doc = Document(
                        'Templates/Максимальные зоны смертельного поражения при реализации аварий со взрывом ТВС.docx')  ####################

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]

                    # if var == 'RPR':
                    #     row = doc_table[2].add_row().cells  # олучаем все ячейки ряда
                    # elif var == 'OPO':
                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда

                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/1.docx')
                print("--- %s seconds ---" % (time.time() - start_time))

                self._signal.emit(int(percent))

            elif 'Огненный шар Вероятностно' in sheet_name[i]:  # i = 4

                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []

                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    if not (obor in oborudovanie): oborudovanie.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)
                    s_def_otv = str(sheet['F' + str(j)].value)  # Площадь деф отв)

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)
                    if not (ve in ov): ov.append(ve)

                    stroka.append(s_def_otv)
                    if not (s_def_otv in pdo): pdo.append(s_def_otv)
                    stroka.append(sheet['G' + str(j)].value)  # Масса
                    stroka.append(sheet['I' + str(j)].value)  # Радиус 1
                    stroka.append(sheet['J' + str(j)].value)  # Радиус 10
                    stroka.append(sheet['K' + str(j)].value)  # Радиус 25
                    stroka.append(sheet['L' + str(j)].value)  # Радиус 50
                    stroka.append(sheet['M' + str(j)].value)  # Радиус 90
                    stroka.append(sheet['N' + str(j)].value)  # Радиус 99
                    stroka.append(sheet['O' + str(j)].value)  # Радиус 100

                    dict.append(stroka)

                dict.sort(key=lambda row: (row[0] or 0))
                # print(dict)
                rez = chistka_3_ogn(dict)

                count_stolb = 10

                if var == 'RPR':
                    doc = Document(
                        'Templates/Результаты расчета зон действия поражающих факторов при огненном шаре.docx')  ####################
                elif var == 'OPO':
                    doc = Document(
                        'Templates/Максимальные зоны смертельного поражения при реализации аварий с «огненным шаром».docx')  ####################

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]

                    # if var == 'RPR':
                    #     row = doc_table[4].add_row().cells  # олучаем все ячейки ряда
                    # elif var == 'OPO':
                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда

                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/2.docx')
                print("--- %s seconds ---" % (time.time() - start_time))

                self._signal.emit(int(percent))

            elif 'Горизонтальный факел Веро' in sheet_name[i]:  # i = 3

                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []
                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)  # Вещество
                    s_def_otv = str(sheet['F' + str(j)].value)  # Площадь деф отв

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)

                    stroka.append(s_def_otv)  # Площадь деф отв

                    stroka.append(sheet['G' + str(j)].value)  # Масса
                    stroka.append(sheet['I' + str(j)].value)  # Радуис
                    dict.append(stroka)

                dict.sort(key=lambda row: (row[2] or 0, row[3] or 0, row[0] or 0))
                # print(dict)
                rez = chistka_3_ogn(dict)

                count_stolb = 4

                if var == 'RPR':
                    doc = Document(
                        'Templates/Результаты расчета зон действия поражающих факторов при струйном горении (горизонтальный факел) РПР.docx')  ####################
                elif var == 'OPO':
                    doc = Document(
                        'Templates/Максимальные зоны смертельного поражения при реализации аварий со струйным горением (горизонтальный факел).docx')  ####################

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]

                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда

                    # if var == 'RPR':
                    #     row = doc_table[6].add_row().cells  # олучаем все ячейки ряда
                    # elif var == 'OPO':
                    #     row = doc_table[7].add_row().cells  # олучаем все ячейки ряда

                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/3.docx')
                print("--- %s seconds ---" % (time.time() - start_time))

                self._signal.emit(int(percent))
            elif 'Вертикальный факел Вероят' in sheet_name[i]:  # i = 1

                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []

                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    if not (obor in oborudovanie): oborudovanie.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)
                    s_def_otv = str(sheet['F' + str(j)].value)  # Площадь деф отв)

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)
                    if not (ve in ov): ov.append(ve)

                    stroka.append(s_def_otv)
                    if not (s_def_otv in pdo): pdo.append(s_def_otv)
                    stroka.append(sheet['G' + str(j)].value)  # Масса
                    stroka.append(sheet['I' + str(j)].value)  # Радиус 1
                    stroka.append(sheet['J' + str(j)].value)  # Радиус 10
                    stroka.append(sheet['K' + str(j)].value)  # Радиус 25
                    stroka.append(sheet['L' + str(j)].value)  # Радиус 50
                    stroka.append(sheet['M' + str(j)].value)  # Радиус 90
                    stroka.append(sheet['N' + str(j)].value)  # Радиус 99
                    stroka.append(sheet['O' + str(j)].value)  # Радиус 100

                    dict.append(stroka)

                dict.sort(key=lambda row: (row[2] or 0, row[3] or 0, row[0] or 0))
                # print(dict)
                rez = filtery_for_1(dict, oborudovanie, ov, pdo, 5)

                count_stolb = 10

                if var == 'RPR':
                    doc = Document(
                        'Templates/Результаты расчета зон действия поражающих факторов при струйном горении (вертикальный факел) РПР.docx')  ####################
                elif var == 'OPO':
                    doc = Document(
                        'Templates/Максимальные зоны смертельного поражения при реализации аварий со струйным горением (вертикальный факел).docx')  ####################

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]

                    # if var == 'RPR':
                    #     row = doc_table[10].add_row().cells  # олучаем все ячейки ряда
                    # elif var == 'OPO':
                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда

                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/4.docx')
                print("--- %s seconds ---" % (time.time() - start_time))
                self._signal.emit(int(percent))

            elif 'Пожар-вспышка (методика М' in sheet_name[i]:  # RPR i = 0

                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []
                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    if not (obor in oborudovanie): oborudovanie.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)  # Вещество
                    s_def_otv = str(sheet['H' + str(j)].value)  # Площадь деф отв

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)
                    if not (ve in ov): ov.append(ve)

                    stroka.append(s_def_otv)  # Площадь деф отв
                    if not (s_def_otv in pdo): pdo.append(s_def_otv)

                    stroka.append(sheet['I' + str(j)].value)  # Масса
                    stroka.append(sheet['K' + str(j)].value)  # Радиус НКПВ
                    dict.append(stroka)

                dict.sort(key=lambda row: (row[2] or 0, row[3] or 0, row[0] or 0))
                # print(dict)
                rez = filtery_for_1(dict, oborudovanie, ov, pdo, 5)

                count_stolb = 4

                doc = Document(
                    'Templates/Результаты расчета зон действия поражающих факторов при пожаре-вспышке.docx')

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]

                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда

                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/5.docx')
                print("--- %s seconds ---" % (time.time() - start_time))
                print("----------------------------")
                print(len(sheet_name))

                print(percent)
                self._signal.emit(int(percent))
                print("----------------------------")

            # OPO!!!!!!

            elif 'Пожар-вспышка по РБ РТН' in sheet_name[i]:  # i = 7

                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []

                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    if not (obor in oborudovanie): oborudovanie.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)

                    s_def_otv = str(sheet['H' + str(j)].value)  # Площадь деф отв

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)
                    if not (ve in ov): ov.append(ve)

                    # s = sheet['H' + str(j)].value # Площадь деф отв
                    stroka.append(s_def_otv)
                    if not (s_def_otv in pdo): pdo.append(s_def_otv)

                    stroka.append(round(sheet['I' + str(j)].value, 2))  # Масса
                    stroka.append(round(sheet['J' + str(j)].value, 2))  # Дрейф
                    stroka.append(sheet['L' + str(j)].value)  # По ветру
                    stroka.append(sheet['M' + str(j)].value)  # Против ветра
                    stroka.append(sheet['N' + str(j)].value)  # Полуширина
                    dict.append(stroka)

                dict.sort(key=lambda row: ((row[0]) or 0, row[2] or 0, row[3] or 0))
                # print(dict)
                # Максимальная масса, полуширина, дрейф
                rez = filtery_for_3(dict, oborudovanie, ov, pdo, 4, 8, 5)

                count_stolb = 7

                doc = Document(
                    'Templates/Максимальные зоны смертельного поражения при реализации аварий с «пожаром-вспышкой».docx')

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]

                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда

                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/6.docx')
                print("--- %s seconds ---" % (time.time() - start_time))

                self._signal.emit(int(percent))

            elif 'Взрыв ТВС Избыточное давл' in sheet_name[i] and var == 'RPR':
                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []

                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    if not (obor in oborudovanie): oborudovanie.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)
                    s_def_otv = str(sheet['H' + str(j)].value)  # Площадь деф отв

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)
                    if not (ve in ov): ov.append(ve)

                    stroka.append(s_def_otv)
                    if not (s_def_otv in pdo): pdo.append(s_def_otv)

                    stroka.append(sheet['I' + str(j)].value)  # Масса
                    stroka.append(sheet['J' + str(j)].value)  # Дрейф
                    stroka.append(sheet['K' + str(j)].value)  # Радиус 3
                    stroka.append(sheet['L' + str(j)].value)  # Радиус 5
                    stroka.append(sheet['M' + str(j)].value)  # Радиус 12
                    stroka.append(sheet['N' + str(j)].value)  # Радиус 28
                    stroka.append(sheet['O' + str(j)].value)  # Радиус 53
                    stroka.append(sheet['P' + str(j)].value)  # Радиус 100
                    dict.append(stroka)
                # print(dict)  # [Оборудование, Состояние, Вещество, ПДО, Масса, Дрейф, Радиус]
                dict.sort(key=lambda row: ((row[2]) or 0, row[3] or 0, row[0] or 0))
                # print(dict)

                rez = filtery(dict, oborudovanie, ov, pdo, 6, 5)
                # print(rez)  # Оборудование, Состояние, Масса, Дрейф, Радиус
                # rez = chistka_for_pozh_intens(rez)
                for i in range(len(rez)):
                    rez[i].pop(3)
                count_stolb = 9

                doc = Document(
                    'Templates/Результаты расчета зон действия поражающих факторов при дефлаграционном горении облака ТВС РПР.docx')

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]
                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда
                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/7.docx')
                print("--- %s seconds ---" % (time.time() - start_time))

                self._signal.emit(int(percent))

            elif 'Взрыв ТВС Избыточное давл' in sheet_name[i] and var == 'OPO':
                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []

                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    if not (obor in oborudovanie): oborudovanie.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)
                    s_def_otv = str(sheet['H' + str(j)].value)  # Площадь деф отв

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)
                    if not (ve in ov): ov.append(ve)

                    stroka.append(s_def_otv)
                    if not (s_def_otv in pdo): pdo.append(s_def_otv)

                    stroka.append(sheet['I' + str(j)].value)  # Масса
                    stroka.append(sheet['J' + str(j)].value)  # Дрейф
                    stroka.append(sheet['K' + str(j)].value)  # Радиус 2
                    stroka.append(sheet['L' + str(j)].value)  # Радиус 3
                    stroka.append(sheet['M' + str(j)].value)  # Радиус 5
                    stroka.append(sheet['N' + str(j)].value)  # Радиус 12
                    stroka.append(sheet['O' + str(j)].value)  # Радиус 14
                    stroka.append(sheet['P' + str(j)].value)  # Радиус 28
                    stroka.append(sheet['Q' + str(j)].value)  # Радиус 53
                    stroka.append(sheet['R' + str(j)].value)  # Радиус 70
                    stroka.append(sheet['S' + str(j)].value)  # Радиус 100
                    dict.append(stroka)
                print(dict)  # [Оборудование, Состояние, Вещество, ПДО, Масса, Дрейф, Радиус]
                dict.sort(key=lambda row: ((row[2]) or 0, row[3] or 0, row[0] or 0))
                # print(dict)

                rez = filtery(dict, oborudovanie, ov, pdo, 6, 5)
                print(rez)  # Оборудование, Состояние, Масса, Дрейф, Радиус
                for i in range(len(rez)):
                    rez[i].pop(3)
                count_stolb = 12
                # rez = chistka_for_pozh_intens(rez)
                doc = Document(
                    'Templates/Результаты расчета зон действия поражающих факторов при дефлаграционном горении облака ТВС.docx')

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]
                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда
                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/12.docx')
                print("--- %s seconds ---" % (time.time() - start_time))

                self._signal.emit(int(percent))

            elif 'Пожар пролива Интенсивнос' in sheet_name[i]:
                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []
                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    if not (obor in oborudovanie): oborudovanie.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)  # Вещество
                    s_def_otv = str(sheet['H' + str(j)].value)  # Площадь деф отв

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)
                    if not (ve in ov): ov.append(ve)

                    stroka.append(s_def_otv)  # Площадь деф отв
                    if not (s_def_otv in pdo): pdo.append(s_def_otv)

                    stroka.append(sheet['G' + str(j)].value)  # Площадь пролива

                    stroka.append(sheet['J' + str(j)].value)  # Дрейф
                    stroka.append(sheet['K' + str(j)].value)  # Радиус 1,4
                    stroka.append(sheet['L' + str(j)].value)  # Радиус 4,2
                    stroka.append(sheet['M' + str(j)].value)  # Радиус 7
                    stroka.append(sheet['N' + str(j)].value)  # Радиус 10,5
                    stroka.append(sheet['O' + str(j)].value)  # Радиус 13,9
                    stroka.append(sheet['P' + str(j)].value)  # Радиус 14,8
                    dict.append(stroka)

                # print(dict) #[Оборудование, Состояние, Вещество, ПДО, Площадь пролива, дрейф, радиус1 ]
                dict.sort(key=lambda row: (row[2] or 0, row[3] or 0, row[0] or 0))
                # print(dict)

                rez = filtery(dict, oborudovanie, ov, pdo, 6, 5)

                # print(rez) #[Оборудование, Состояние, Площадь пролива, дрейф, радиус1 ]

                print(rez)

                count_stolb = 9

                doc = Document(
                    'Templates/Результаты расчета зон действия поражающих факторов при пожаре-пролива.docx')  ####################

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]
                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда
                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/8.docx')

                print("--- %s seconds ---" % (time.time() - start_time))

                self._signal.emit(int(percent))

            elif 'Огненный шар Интенсивност' in sheet_name[i]:
                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []
                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)  # Вещество
                    s_def_otv = str(sheet['F' + str(j)].value)  # Площадь деф отв

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(sheet['G' + str(j)].value)  # Масса

                    stroka.append(sheet['I' + str(j)].value)  # Радиус 1,4
                    stroka.append(sheet['J' + str(j)].value)  # Радиус 4,2
                    stroka.append(sheet['K' + str(j)].value)  # Радиус 7
                    stroka.append(sheet['L' + str(j)].value)  # Радиус 10,5
                    stroka.append(sheet['M' + str(j)].value)  # Радиус 13,9
                    stroka.append(sheet['N' + str(j)].value)  # Радиус 14,8

                    dict.append(stroka)

                # print(dict) #[Оборудование, Масса, радиус1 ]
                dict.sort(key=lambda row: (row[0] or 0))

                rez = dict

                rez = [x for x in rez if x != []]

                count_stolb = 9

                doc = Document(
                    'Templates/Максимальные зоны смертельного поражения при реализации аварий с «огненным шаром» РПР.docx')

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]
                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда
                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/9.docx')

                self._signal.emit(int(percent))

            elif 'Горизонтальный факел Инте' in sheet_name[i]:  # Готово
                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []
                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)  # Вещество
                    stroka.append(ve)

                    stroka.append(sheet['H' + str(j)].value)  # Расход
                    stroka.append(sheet['I' + str(j)].value)  # Радиус 10
                    stroka.append(sheet['J' + str(j)].value)  # Радиус 100
                    dict.append(stroka)

                # print(dict) #[Оборудование, Вещество, Расход, радиус1 ]
                dict.sort(key=lambda row: (row[1] or 0, row[0]))
                # print(dict)

                # print(dict)

                rez = dict
                rez = [x for x in rez if x != []]

                count_stolb = 5

                doc = Document(
                    'Templates/Результаты расчета зон действия поражающих факторов при струйном горении (горизонтальный факел).docx')

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]
                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда
                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/10.docx')
                print("--- %s seconds ---" % (time.time() - start_time))

                self._signal.emit(int(percent))

            elif 'Вертикальный факел Интенс' in sheet_name[i]:  #
                dict = []
                l = 0
                print(sheet_name[i])
                sheet = wb.get_sheet_by_name(sheet_name[i])
                all_rows = sheet.max_row
                oborudovanie = []
                ov = []
                pdo = []

                for j in range(4, all_rows):
                    stroka = []
                    obor = str(sheet['C' + str(j)].value).strip()  # Оборудование
                    stroka.append(obor)
                    if not (obor in oborudovanie): oborudovanie.append(obor)

                    ve = str(sheet['E' + str(j)].value).strip()  # Вещество
                    ve = change_gas(ve)  # Вещество
                    s_def_otv = str(sheet['F' + str(j)].value)  # Площадь деф отв

                    sostoyan = str(sheet['D' + str(j)].value).strip()  # Cостояние
                    state = change_state_with_old(s_def_otv, ve, sostoyan)  # Состояние
                    stroka.append(state)

                    stroka.append(ve)
                    if not (ve in ov): ov.append(ve)

                    stroka.append(s_def_otv)  # Площадь деф отв
                    if not (s_def_otv in pdo): pdo.append(s_def_otv)

                    stroka.append(sheet['H' + str(j)].value)  # Расход

                    stroka.append(sheet['I' + str(j)].value)  # Радиус 1,4
                    stroka.append(sheet['J' + str(j)].value)  # Радиус 4,2
                    stroka.append(sheet['K' + str(j)].value)  # Радиус 7
                    stroka.append(sheet['L' + str(j)].value)  # Радиус 10,5
                    stroka.append(sheet['M' + str(j)].value)  # Радиус 13,9
                    stroka.append(sheet['N' + str(j)].value)  # Радиус 14,8
                    dict.append(stroka)

                # print(dict) #[Оборудование, Состояние, Вещество, ПДО, Расход, радиус1, ]
                dict.sort(key=lambda row: (row[2] or 0, row[3] or 0, row[0] or 0))

                # print(dict)

                rez = filtery_for_vert_li_fakel(dict, oborudovanie, ov, pdo, 5)
                # print(rez)
                rez = chistka_vert_fakel(rez)

                # print(rez) #[Оборудование, Состояние, Площадь пролива, дрейф, радиус1 ]
                # rez = chistka_for_pozh_intens(rez)
                # print(rez)

                count_stolb = 9

                doc = Document(
                    'Templates/Результаты расчета зон действия поражающих факторов при струйном горении (вертикальный факел).docx')  ####################

                doc_table = doc.tables
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(10)
                paragraph_format = doc.styles['Normal'].paragraph_format
                paragraph_format.line_spacing = Pt(12)
                paragraph_format.space_after = (0)

                for k in range(len(rez)):
                    list = rez[k]
                    row = doc_table[0].add_row().cells  # олучаем все ячейки ряда
                    for col in range(count_stolb):
                        cell = row[col]
                        paragraph = cell.paragraphs[0]
                        paragraph.style = doc.styles['Normal']
                        if col == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if str(list[col]) is None or str(list[col]) == 'None': list[col] = ''
                        run = paragraph.add_run(str(list[col]))
                doc.save('output/11.docx')
                print("--- %s seconds ---" % (time.time() - start_time))

                self._signal.emit(int(percent))


        start_combine(self.word_path)

        self._signal.emit(100)
        time.sleep(2)
        self.finished.emit(0)

    # def terminate(self):
    #     print("terminate func enter")
    #     self.finished.emit(2)
    #     return 123


